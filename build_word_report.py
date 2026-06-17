"""Build the final DOCX experiment report."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RESULT_DIR = ROOT / "cnn_output" / "domain_adaptation"
FIG_DIR = ROOT / "reports" / "figures"
OUT_PATH = ROOT / "reports" / "PPG_BP_小样本重建_领域自适应与LwF实验报告.docx"


def load_results() -> dict[int, dict]:
    return {
        shot: json.load(open(RESULT_DIR / f"fewshot_{shot}_results.json", "r", encoding="utf-8"))
        for shot in (16, 32, 64)
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 10),
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("PPG-BP 小样本重建实验报告")
    set_run_font(run, 9, color="666666")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, 9.5, color="555555")


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r, 10.5)


def add_result_table(doc: Document, results: dict[int, dict], method: str) -> None:
    rows = 1 + len(results)
    table = doc.add_table(rows=rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["样本数", "SBP MAE", "SBP STD", "DBP MAE", "DBP STD"]
    for i, header in enumerate(headers):
        set_cell_shading(table.cell(0, i), "F2F4F7")
        set_cell_text(table.cell(0, i), header, bold=True)

    for row_idx, shot in enumerate((16, 32, 64), start=1):
        r = results[shot][method]
        values = [
            f"{shot}-shot",
            f"{r['SBP_MAE']:.2f}",
            f"{r['SBP_STD']:.2f}",
            f"{r['DBP_MAE']:.2f}",
            f"{r['DBP_STD']:.2f}",
        ]
        for col_idx, val in enumerate(values):
            set_cell_text(table.cell(row_idx, col_idx), val)

    doc.add_paragraph()


def add_baseline_table(doc: Document, results: dict[int, dict]) -> None:
    table = doc.add_table(rows=2, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["模型", "SBP MAE", "SBP STD", "DBP MAE", "DBP STD"]
    for i, header in enumerate(headers):
        set_cell_shading(table.cell(0, i), "F2F4F7")
        set_cell_text(table.cell(0, i), header, bold=True)
    base = results[16]["base"]
    values = ["Base", f"{base['SBP_MAE']:.2f}", f"{base['SBP_STD']:.2f}", f"{base['DBP_MAE']:.2f}", f"{base['DBP_STD']:.2f}"]
    for i, val in enumerate(values):
        set_cell_text(table.cell(1, i), val)
    doc.add_paragraph()


def add_method_section(doc: Document, title: str, method: str, results: dict[int, dict]) -> None:
    doc.add_heading(title, level=1)
    if method == "feature_da":
        add_para(
            doc,
            "该方法从 base 模型初始化 student，冻结浅层特征提取模块，在目标域小样本监督损失之外加入源域与目标域 GAP 特征的 CORAL 对齐损失。它直接约束中间特征分布，是本实验中收益最稳定的方法。",
        )
    else:
        add_para(
            doc,
            "该方法复制 base 模型作为冻结 teacher。student 在目标域小样本上学习标签，同时在源域 batch 上保持 teacher 输出，以减少迁移过程中的旧知识遗忘。",
        )

    doc.add_heading("结果表", level=2)
    add_result_table(doc, results, method)

    doc.add_heading("指标图", level=2)
    doc.add_picture(str(FIG_DIR / f"{method}_metric_bars.png"), width=Inches(6.4))
    add_caption(doc, f"图：{title} 在不同小样本规模下的 MAE/STD 指标")

    doc.add_picture(str(FIG_DIR / f"{method}_mae_trends.png"), width=Inches(6.2))
    add_caption(doc, f"图：{title} 的 SBP/DBP MAE 随样本量变化趋势")

    doc.add_picture(str(FIG_DIR / f"{method}_64_scatter.png"), width=Inches(6.2))
    add_caption(doc, f"图：64-shot 下 {title} 与 Base 的真实值-预测值散点图")


def build_docx() -> Path:
    results = load_results()
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PPG-BP 小样本重建实验报告")
    set_run_font(run, 22, True, "0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("特征级领域自适应与 LwF 方法对比")
    set_run_font(run, 12, color="555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("实验日期：2026-06-17    目标域：自建 PPG 数据集    源域：公开 PPG 数据集")
    set_run_font(run, 10, color="666666")

    doc.add_heading("摘要", level=1)
    add_para(
        doc,
        "本报告基于已有 1D-ResNet 血压预测模型，在自建 PPG 目标域上进行小样本重建实验。实验分别评估特征级领域自适应和 LwF 两种方法，并在 16-shot、32-shot、64-shot 设置下报告 SBP/DBP 的 MAE 与误差标准差。",
    )
    add_bullets(
        doc,
        [
            "64-shot 下，特征级领域自适应取得最佳结果：SBP MAE 7.69 mmHg，DBP MAE 2.90 mmHg。",
            "相对 Base，64-shot 特征级领域自适应使 SBP MAE 降低 32.36%，DBP MAE 降低 59.57%。",
            "LwF 同样有效，但在本实验设置下整体略逊于特征级领域自适应。",
        ],
    )

    doc.add_heading("数据与实验设置", level=1)
    add_para(doc, "源域公开数据来自已解压的数据集目录，训练集包含 4745 条 PPG 波形及 SBP/DBP 标签。目标域自建数据来自外部 dataset 目录，训练集 253 条，测试集 152 条。")
    add_para(doc, "所有实验均使用已有 base 模型权重 models/base_model_best.pth 初始化。PPG 输入被处理为 PPG、一阶导数、二阶导数三通道，并使用 models/norm_params.npz 中的归一化参数。")
    add_bullets(
        doc,
        [
            "小样本规模：16-shot、32-shot、64-shot。",
            "训练参数：epochs=40，batch size=16，learning rate=1e-4，optimizer=Adam，gradient clip=5.0。",
            "评价指标：SBP/DBP 的 MAE 和预测误差 STD，单位为 mmHg。",
        ],
    )

    doc.add_heading("Base 跨域基线", level=1)
    add_para(doc, "Base 模型未经目标域适配，直接在自建测试集上评估，作为两种小样本方法的对照。")
    add_baseline_table(doc, results)

    add_method_section(doc, "特征级领域自适应", "feature_da", results)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_method_section(doc, "LwF", "lwf", results)

    doc.add_heading("结论", level=1)
    add_para(
        doc,
        "从实验结果看，目标域样本达到 32 条后，两种方法都能明显改善 SBP 和 DBP 预测误差；样本数增加到 64 条后，性能进一步提升。特征级领域自适应在 32-shot 和 64-shot 下均优于 LwF，说明当前公开域与自建域之间存在显著中间特征分布偏移，直接对齐特征分布比仅保持 teacher 输出更适合本任务。",
    )
    add_para(
        doc,
        "因此，在当前数据和超参数设置下，推荐优先采用特征级领域自适应作为小样本重建方法；LwF 可作为更保守的适配方案，在强调保留源域模型行为时使用。",
    )

    doc.add_heading("复现实验命令", level=1)
    cmd = (
        ".\\.venv\\Scripts\\python.exe domain_adaptation_experiments.py "
        "--public-dir \"D:\\专业实习\\ppg-bp-prediction-master\\_provided_dataset_run\\数据集\\1、公开数据集\" "
        "--self-dir \"D:\\专业实习\\程序 - 5.13(1)\\程序 - 5.13\\dataset\" "
        "--shots 64 --epochs 40 --batch-size 16 --cpu"
    )
    p = doc.add_paragraph()
    r = p.add_run(cmd)
    set_run_font(r, 9, color="333333")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_docx())

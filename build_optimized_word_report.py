from pathlib import Path
import json

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
RESULT_DIR = ROOT / "cnn_output" / "domain_adaptation"
FIG_DIR = ROOT / "reports" / "figures_optimized"
OUT = ROOT / "reports" / "PPG_BP_小样本重建_优化后实验报告.docx"
SHOTS = [64, 128, 253]


def load_results():
    return {s: json.load(open(RESULT_DIR / f"fewshot_{s}_calibrated_results.json", encoding="utf-8")) for s in SHOTS}


def font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell(cell_, text, bold=False):
    cell_.text = ""
    p = cell_.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, 9.5, bold)
    cell_.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.1
    for name, size, color in [("Title", 22, "0B2545"), ("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        font(r, 10.5)


def result_table(doc, results, method):
    table = doc.add_table(rows=1 + len(SHOTS), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["样本数", "SBP MAE", "SBP STD", "DBP MAE", "DBP STD"]
    for i, h in enumerate(headers):
        shade(table.cell(0, i), "F2F4F7")
        cell(table.cell(0, i), h, True)
    for r_i, shot in enumerate(SHOTS, 1):
        row = results[shot][method]
        vals = [f"{shot}-shot", f"{row['SBP_MAE']:.2f}", f"{row['SBP_STD']:.2f}", f"{row['DBP_MAE']:.2f}", f"{row['DBP_STD']:.2f}"]
        for c_i, val in enumerate(vals):
            cell(table.cell(r_i, c_i), val)
    doc.add_paragraph()


def baseline_table(doc, results):
    table = doc.add_table(rows=1 + len(SHOTS), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["校准样本数", "SBP MAE", "SBP STD", "DBP MAE", "DBP STD"]):
        shade(table.cell(0, i), "F2F4F7")
        cell(table.cell(0, i), h, True)
    for r_i, shot in enumerate(SHOTS, 1):
        row = results[shot]["base"]
        vals = [f"{shot}-shot", f"{row['SBP_MAE']:.2f}", f"{row['SBP_STD']:.2f}", f"{row['DBP_MAE']:.2f}", f"{row['DBP_STD']:.2f}"]
        for c_i, val in enumerate(vals):
            cell(table.cell(r_i, c_i), val)
    doc.add_paragraph()


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    font(r, 9.5, color="555555")


def method_section(doc, title, method, results):
    doc.add_heading(title, level=1)
    if method == "feature_da":
        para(doc, "优化后的特征级领域自适应在原 CORAL 特征对齐基础上加入目标域输出校准。该校准仅使用目标域训练小样本拟合预测值到真实血压的仿射映射，用于修正系统性偏移。")
    else:
        para(doc, "优化后的 LwF 在源域 teacher 输出保持约束基础上加入同样的目标域输出校准。该方法更保守，但在全量目标域训练样本下 SBP 表现最好。")
    doc.add_heading("优化后结果", level=2)
    result_table(doc, results, method)
    doc.add_picture(str(FIG_DIR / f"{method}_optimized_bars.png"), width=Inches(6.4))
    caption(doc, f"{title}：优化后指标柱状图")
    doc.add_picture(str(FIG_DIR / f"{method}_optimized_trend.png"), width=Inches(6.2))
    caption(doc, f"{title}：优化后 MAE 趋势图")
    doc.add_picture(str(FIG_DIR / f"{method}_optimized_scatter.png"), width=Inches(6.2))
    caption(doc, f"{title}：253-shot 优化后散点图")


def build():
    results = load_results()
    doc = Document()
    setup(doc)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PPG-BP 小样本重建优化后实验报告")
    font(r, 22, True, "0B2545")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("特征级领域自适应与 LwF 独立优化结果")
    font(r, 12, color="666666")

    doc.add_heading("摘要", level=1)
    para(doc, "针对初版散点图预测分散的问题，本次优化在两种迁移方法后加入目标域小样本输出校准，并将训练轮数提高到 60。所有结果均由真实自建测试集计算得到。")
    bullets(doc, [
        "特征级领域自适应最佳结果：253-shot，SBP MAE 4.62 mmHg，DBP MAE 2.17 mmHg。",
        "LwF 最佳结果：253-shot，SBP MAE 4.36 mmHg，DBP MAE 2.34 mmHg。",
        "两种方法的图表已分开绘制，分别包含指标柱状图、MAE 趋势图和散点图。",
    ])

    doc.add_heading("优化策略", level=1)
    bullets(doc, [
        "保留原有 1D-ResNet 预训练权重作为初始化。",
        "特征级领域自适应继续使用 CORAL 中间特征分布对齐。",
        "LwF 继续使用 teacher 输出保持约束。",
        "新增目标域输出校准：使用目标域训练小样本拟合 ridge 仿射校准层，纠正跨域系统偏移。",
        "训练设置：epochs=60，batch size=16，learning rate=1e-4，CPU 运行。",
    ])

    doc.add_heading("Base + 校准基线", level=1)
    para(doc, "Base 模型不做迁移训练，仅使用目标域小样本进行输出校准，作为优化后方法的基线。")
    baseline_table(doc, results)

    method_section(doc, "特征级领域自适应", "feature_da", results)
    doc.add_page_break()
    method_section(doc, "LwF", "lwf", results)

    doc.add_heading("结论", level=1)
    para(doc, "加入目标域输出校准后，两种方法的散点图和 MAE 指标均明显改善。特征级领域自适应在 DBP 上表现最好，LwF 在 253-shot 下获得最低 SBP MAE。综合 SBP 与 DBP，本实验推荐优先采用特征级领域自适应；若更重视 SBP，可选择 LwF 优化版。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
